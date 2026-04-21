"""Generate a beautifully formatted Word document from Project_description.md"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color Palette ────────────────────────────────────────────
C_NAVY    = RGBColor(0x1A, 0x2B, 0x4A)   # title / H1
C_BLUE    = RGBColor(0x1F, 0x5C, 0x99)   # H2
C_TEAL    = RGBColor(0x17, 0x7A, 0x7E)   # H3
C_ACCENT  = RGBColor(0xE8, 0x6D, 0x2A)   # inline highlight
C_GRAY    = RGBColor(0x55, 0x55, 0x55)   # body text
C_LGRAY   = RGBColor(0xF2, 0xF4, 0xF8)   # code / table header bg
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_BDR     = RGBColor(0xCC, 0xD5, 0xE0)   # table border

FONT_BODY = "Noto Sans TC"
FONT_CODE = "Consolas"
FONT_HEAD = "Noto Sans TC"

# ── Helpers ──────────────────────────────────────────────────

def hex_to_str(c: RGBColor) -> str:
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_to_str(color))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), hex_to_str(C_BDR))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)


def add_run(para, text, bold=False, italic=False, color=None,
            font_name=None, size=None, underline=False):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.name = font_name or FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name or FONT_BODY)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_heading1(doc, text):
    p = doc.add_paragraph()
    para_spacing(p, before=24, after=8)
    p.paragraph_format.keep_with_next = True
    # left border decoration
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), hex_to_str(C_BLUE))
    left.set(qn("w:space"), "6")
    pBdr.append(left)
    pPr.append(pBdr)
    add_run(p, text, bold=True, color=C_NAVY, font_name=FONT_HEAD, size=16)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    para_spacing(p, before=14, after=4)
    p.paragraph_format.keep_with_next = True
    add_run(p, "▍ ", bold=True, color=C_TEAL, font_name=FONT_HEAD, size=13)
    add_run(p, text, bold=True, color=C_BLUE, font_name=FONT_HEAD, size=13)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    para_spacing(p, before=10, after=3)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, bold=True, color=C_TEAL, font_name=FONT_HEAD, size=11.5)
    return p


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    para_spacing(p, before=2, after=4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, text, color=C_GRAY, size=10.5)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.6 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    para_spacing(p, before=1, after=1)
    # parse **bold** inline
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        add_run(p, part, bold=(i % 2 == 1), color=C_GRAY, size=10.5)
    return p


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    para_spacing(p, before=4, after=4)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    # shaded background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_to_str(C_LGRAY))
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = FONT_CODE
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CODE)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p


def add_inline_code(para, text):
    run = para.add_run(f" {text} ")
    run.font.name = FONT_CODE
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CODE)
    run.font.size  = Pt(9.5)
    run.font.color.rgb = C_ACCENT
    run.bold = True


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, C_NAVY)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, before=3, after=3)
        run = p.add_run(h)
        run.bold = True
        run.font.name = FONT_HEAD
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
        run.font.size  = Pt(10)
        run.font.color.rgb = C_WHITE

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = C_LGRAY if r_idx % 2 == 0 else C_WHITE
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            para_spacing(p, before=2, after=2)
            # parse **bold**
            parts = cell_text.split("**")
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.bold = (i % 2 == 1)
                run.font.name = FONT_BODY
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
                run.font.size  = Pt(10)
                run.font.color.rgb = C_GRAY

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    para_spacing(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:color"), hex_to_str(C_BDR))
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Document Assembly ─────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Cover Page ───────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(cover, before=60, after=12)
add_run(cover, "基於 HAM 認知模型與多模態大語言模型的\nAI 角色自主生活模擬系統",
        bold=True, color=C_NAVY, font_name=FONT_HEAD, size=22)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(sub, before=0, after=8)
add_run(sub, "專題描述文件", bold=False, color=C_BLUE, font_name=FONT_HEAD, size=14)

sep = doc.add_paragraph()
sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(sep, before=4, after=60)
add_run(sep, "─" * 30, color=C_BDR, size=10)

doc.add_page_break()

# ── 一句話摘要 ────────────────────────────────────────────────
add_heading1(doc, "一句話摘要")
p = doc.add_paragraph()
para_spacing(p, before=4, after=8)
p.paragraph_format.left_indent  = Cm(0.8)
p.paragraph_format.right_indent = Cm(0.8)
# italic quote style
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
for side in ("left",):
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "16")
    el.set(qn("w:color"), hex_to_str(C_ACCENT))
    el.set(qn("w:space"), "8")
    pBdr.append(el)
pPr.append(pBdr)
add_run(p, "本專題建構一套讓多個 AI 角色在 Unreal Engine 虛擬環境中自主生活、互動並形成長期記憶的完整系統，"
           "結合神經科學的人類記憶理論與現代多模態模型，實現具備感知、思考、記憶與遺忘能力的 AI 代理人。",
        italic=True, color=C_GRAY, size=11)

add_divider(doc)

# ── 系統概覽 ──────────────────────────────────────────────────
add_heading1(doc, "系統概覽")
add_body(doc, "本系統讓五個具備獨立個性、職業與人際關係的 AI 角色，在共同的虛擬世界中自主運作。"
              "每個角色透過 Unreal Engine 接收來自虛擬環境的畫面與事件，經由多模態模型做出決策，"
              "並將重要的生活經歷逐漸累積為長期記憶，同時會因時間流逝與使用頻率的差異而產生遺忘，"
              "模擬真實人腦的記憶機制。")
add_divider(doc)

# ── 理論基礎 ──────────────────────────────────────────────────
add_heading1(doc, "理論基礎")

add_heading2(doc, "Human Associative Memory（HAM）模型")
add_body(doc, "Anderson & Bower（1973）提出的 HAM 模型認為，人類長期記憶以「命題」的方式組織，"
              "每個記憶單元由主詞（Subject）、關係（Relation）、受詞（Object）三個節點構成，"
              "並可附加地點與時間資訊。本系統以此結構作為角色長期記憶（LTM）的儲存格式：")
add_code_block(doc, "示例命題：Amy 遇見 David（咖啡廳，早上）\n"
                    "主詞: Amy  |  關係: 遇見  |  受詞: David  |  地點: 咖啡廳  |  時間: 早上")
add_body(doc, "每筆命題都帶有 strength（強度）值，反映該記憶的重要程度與鮮明程度。")

add_heading2(doc, "Atkinson-Shiffrin 多重記憶模型")
add_body(doc, "記憶分為兩層：")
add_bullet(doc, "**STM（短期記憶）**：每天最多 15 筆，記錄當天的場景、事件、決策")
add_bullet(doc, "**LTM（長期記憶）**：從 STM 精選後長期保存，具有衰減與遺忘機制")
add_body(doc, "每天角色「睡覺」時，由 AI 模型判斷今日 STM 中哪些事件值得永久記憶，"
              "模擬人腦在睡眠期間進行的記憶鞏固過程（Memory Consolidation）。")

add_heading2(doc, "Ebbinghaus 遺忘曲線與主動遺忘")
add_body(doc, "LTM 中每筆命題的強度每天衰減，但越常被「想起來」的記憶衰減越慢：")
add_code_block(doc, "actual_decay = DECAY_RATE / (1 + access_count × 0.5)\n\n"
                    "access_count = 0（從未被想起）  → decay = 0.05（快速遺忘）\n"
                    "access_count = 2（被想起兩次）  → decay = 0.025（慢速遺忘）\n"
                    "access_count = 10（常被想起）   → decay = 0.008（幾乎不忘）")
add_body(doc, "強度低於 0.2 的記憶會被自動刪除，模擬「徹底遺忘」。")

add_heading2(doc, "Expected Value of Control（EVC）困惑指數")
add_body(doc, "參考 Shenhav、Botvinick & Cohen（2013）的前扣帶皮層（ACC）功能理論。"
              "人腦會根據當前情境的複雜程度，動態分配認知資源——簡單情境走直覺，複雜情境才深思。")
add_body(doc, "本系統計算困惑指數 C = w₁U + w₂K + w₃S：")

add_styled_table(doc,
    headers=["分量", "意義", "計算方式"],
    rows=[
        ["**U**（不確定性）", "角色不知道該做什麼", "根據候選行動的評分差距計算"],
        ["**K**（衝突）",    "場景資訊相互矛盾",   "關鍵字對衝突偵測（如「陌生人」vs「認識的人」）"],
        ["**S**（驚訝）",    "意料之外的事發生",   "場景新穎度 + 行動停滯懲罰"],
    ],
    col_widths=[3.5, 4.5, 7.5],
)
add_body(doc, "C 值超過閾值 → 深度思考路徑（deliberate）；低於閾值 → 直覺快速決策（intuitive）。")
add_divider(doc)

# ── 五個角色設計 ───────────────────────────────────────────────
add_heading1(doc, "五個角色設計")

add_styled_table(doc,
    headers=["代號", "名字", "職業", "個性關鍵字", "困惑閾值", "個性說明"],
    rows=[
        ["A", "Amy",   "咖啡師",   "溫柔、壓抑、細心",       "0.45（敏感）", "容易陷入深思，對細節很在意"],
        ["B", "Ben",   "超市員工", "開朗、粗神經、不安全感", "0.60（大條）", "大多走直覺，遲鈍不易察覺微妙變化"],
        ["C", "Claire","辦公室員工","直率、謹慎、可靠",       "0.50",        "均衡型，直接表達自己的想法"],
        ["D", "David", "公司老闆", "穩重、內斂、投入工作",   "0.55",        "不輕易改變計畫，對工作高度專注"],
        ["E", "Emma",  "餐廳員工", "活潑、敏感、隨性",       "0.50",        "情緒豐富，容易被氛圍影響"],
    ],
    col_widths=[1.5, 2.0, 3.0, 4.5, 3.2, 5.2],
)

add_heading2(doc, "人際關係網路（設計有戲劇性張力）")
relationships = [
    "Ben → Amy（Ben 追求，Amy 迴避）",
    "Amy ↔ Emma（閨蜜，互相支持）",
    "Claire → Emma（單戀，Emma 尚未察覺）",
    "David → Emma（暗中有好感，不公開）",
    "Claire ↔ David（職場關係，直來直往）",
]
for r in relationships:
    add_bullet(doc, r)

add_divider(doc)

# ── 技術架構 ──────────────────────────────────────────────────
add_heading1(doc, "技術架構")

add_heading2(doc, "整體架構分層")
add_code_block(doc,
"┌─────────────────────────────────────────────────────┐\n"
"│  Unreal Engine（虛擬世界）                            │\n"
"└──────────────────┬──────────────────────────────────┘\n"
"                   │  WebSocket（畫面 + 事件）\n"
"┌──────────────────▼──────────────────────────────────┐\n"
"│  server / perception 層                               │\n"
"│  ws_server.py + yolo_handler.py                      │\n"
"│  接收畫面 → YOLO 偵測 → 轉換為中文語意描述              │\n"
"└──────────────────┬──────────────────────────────────┘\n"
"                   │\n"
"┌──────────────────▼──────────────────────────────────┐\n"
"│  agent 層                                             │\n"
"│  agent_manager.py（協調 5 個角色）                    │\n"
"│  agent.py（單一角色推論流程）                          │\n"
"└──────┬───────────────────────────────────────────────┘\n"
"       │\n"
"  ┌────▼──────────┐        ┌─────────────────────────┐\n"
"  │  core 層       │        │  model 層                │\n"
"  │  character.py │◄──────►│  model_loader.py         │\n"
"  │  stm.py       │        │  text_encoder.py         │\n"
"  │  ltm.py       │        │  vision_encoder.py       │\n"
"  │  confusion.py │        │  fusion_decoder.py       │\n"
"  │  consolidation│        │  prompt_builder.py       │\n"
"  └───────────────┘        │  output_parser.py        │\n"
"                           └─────────────────────────┘\n"
"       │\n"
"┌──────▼──────────────────────────────────────────────┐\n"
"│  config / utils / world 層（全域設定、工具、時鐘）      │\n"
"└─────────────────────────────────────────────────────┘"
)
add_body(doc, "核心設計原則：下層不依賴上層。core/ 完全不碰模型，只透過 callback 接收結果，"
              "這讓核心邏輯可以獨立用 mock 測試，不需要每次都跑 GPU。")
add_divider(doc)

# ── 資料流程 ──────────────────────────────────────────────────
add_heading1(doc, "資料流程（一輪推論）")
add_code_block(doc,
"Unreal Engine 傳入：場景描述 + 畫面 + 事件文字\n"
"         ↓\n"
"YOLO 偵測畫面 → 轉換為中文語意（\"咖啡廳裡有2個人、1個杯子。\"）\n"
"         ↓\n"
"【困惑指數計算】\n"
"  U（不確定）+ K（衝突）+ S（驚訝）→ C 值\n"
"  C < 閾值 → 直覺路徑\n"
"  C ≥ 閾值 → 思考路徑\n"
"         ↓\n"
"【Prompt 組裝】\n"
"  直覺路徑：一句話個性 + 最近 3 筆 STM + LTM 摘要\n"
"  思考路徑：完整個性 + 習慣 + 全部 STM + 相關 LTM + 完整關係描述\n"
"         ↓\n"
"Phi-3.5-Vision 模型推論（本地 GPU）\n"
"         ↓\n"
"【解析輸出】\n"
"  [ACTION] 前往:咖啡廳\n"
"  [THOUGHT] 需要去補貨...\n"
"  [HAM] [\"Amy 前往 咖啡廳（早上）\", ...]\n"
"         ↓\n"
"行動 → WebSocket → Unreal Engine 執行動畫\n"
"HAM 命題 → 寫入 STM\n"
"         ↓\n"
"（角色說出「睡覺」或 STM 滿 15 筆）\n"
"【睡眠濃縮】\n"
"  模型篩選 STM → 重要命題寫入 LTM\n"
"  更新 LTM 摘要、關係描述、今日情緒\n"
"  衰減 LTM strength → 低於 0.2 的命題自動刪除\n"
"  清空 STM → day + 1 → 寫回磁碟"
)
add_divider(doc)

# ── 模組詳解 ──────────────────────────────────────────────────
add_heading1(doc, "模組詳解")

# Config 層
add_heading2(doc, "Config 層（config/）— 全局設定中心")
add_body(doc, "所有可調整的參數集中在此，避免散落在程式各處。")

add_heading3(doc, "world_config.py — 世界規則")
add_styled_table(doc,
    headers=["常數", "值", "用途說明"],
    rows=[
        ["CHARACTER_NAMES", "A→Amy, B→Ben...", "用代碼不用全名，推論時省 token"],
        ["STM_CAPACITY",    "15 筆",             "約等於真實一天的互動量"],
        ["LTM_DECAY_RATE",  "0.05（5%）",        "每天睡眠衰減，模仿遺忘曲線"],
        ["LTM_FORGET_THRESHOLD", "0.2",          "低於此值視為徹底遺忘"],
        ["VALID_EMOTIONS",  "8 種中文情緒",       "限定詞彙，防止模型造出不合格的詞"],
        ["SLEEP_ACTION",    "\"睡覺\"",           "偵測到此動詞就觸發睡眠濃縮流程"],
    ],
    col_widths=[4.5, 4.0, 7.0],
)

add_heading3(doc, "model_config.py — 推論參數")
add_body(doc, "模型：microsoft/Phi-3.5-vision-instruct（4.15B，支援圖片+文字）")
add_styled_table(doc,
    headers=["路徑", "最多 token", "溫度", "說明"],
    rows=[
        ["直覺", "150", "0.0", "快速，節省 GPU；確定性輸出"],
        ["思考", "300", "0.0", "完整推理空間；仍需格式穩定"],
    ],
    col_widths=[3.0, 3.5, 2.5, 7.0],
)
add_body(doc, "溫度設為 0.0（greedy decoding），確保每次相同輸入得到相同格式的輸出，parser 不易出錯。")

add_heading3(doc, "睡眠濃縮 token 預算（RTX 4060 Laptop 效能優化）")
add_styled_table(doc,
    headers=["步驟", "Token 上限", "說明"],
    rows=[
        ["篩選重要命題", "100", "輸出 JSON list，不需長文"],
        ["生成 LTM 摘要", "60",  "1-2 句話"],
        ["更新關係描述", "60",  "1-2 句話"],
        ["推斷今日情緒", "10",  "單一詞彙，如「開心」"],
    ],
    col_widths=[4.5, 3.5, 8.0],
)
add_body(doc, "RTX 4060 Laptop 跑 Phi-3.5 約 1.8 tokens/秒。若每步用預設 256 token，"
              "5 個角色睡眠需要約 47 分鐘；改用以上預算後縮短至 5-8 分鐘。")

add_heading3(doc, "prompts.py — 7 個 Prompt 模板（繁體中文）")
prompts_list = [
    "prompt_intuitive() — 直覺路徑推論",
    "prompt_deliberate() — 思考路徑推論",
    "prompt_select_ltm() — 請模型判斷哪些記憶值得長期保存",
    "prompt_ltm_summary() — 請模型壓縮所有 LTM 為摘要",
    "prompt_update_relationship() — 根據今天對話更新關係描述",
    "prompt_infer_emotion() — 從今天事件推斷當前情緒",
    "prompt_generate_schedule() — 產生排程計畫",
]
for item in prompts_list:
    add_bullet(doc, item)

add_heading3(doc, "action_list.py — 合法行動與地點")
add_bullet(doc, "動詞：前往、回家、對話、工作、休息、吃飯、購物、散步、睡覺等 12 種")
add_bullet(doc, "地點：公寓、咖啡廳、超市、辦公室、餐廳、公園等 16 個")

# Core 層
add_heading2(doc, "Core 層（core/）— 純邏輯，無模型依賴")
add_body(doc, "這層完全不 import 任何模型相關的東西，所有功能都可以用假資料（mock）獨立測試。")

add_heading3(doc, "character.py — 角色狀態容器")
add_code_block(doc,
"靜態（不會改變）：code, name, role, gender, age, personality, habit\n"
"動態（會隨時間更新）：emotion, day, current_location, current_action\n"
"每日清零：today_actions（今天做了什麼）\n\n"
"關係設計有兩層：\n"
"  initial（初始描述）：角色「天生設定」，不可改變\n"
"  summary（動態摘要）：根據實際互動不斷更新，記錄關係的演變"
)

add_heading3(doc, "stm.py — 短期記憶（每筆 turn 結構）")
add_code_block(doc,
"turn_id:          \"D001_T003\"（第1天第3輪）\n"
"scene:            場景描述（\"公園，下午，晴天\"）\n"
"image_desc:       YOLO 語意化結果（\"公園裡有2個人\"）\n"
"input_text:       接收到的對話或事件\n"
"action:           本輪決策行動（\"前往:咖啡廳\"）\n"
"ham_propositions: 本輪抽取的記憶命題 list"
)

add_heading3(doc, "ltm.py — 長期記憶（HAM 命題結構）")
add_code_block(doc,
"id:           \"L001\"（唯一 ID，即使 prune 後也不重複）\n"
"subject:      \"Amy\"\n"
"relation:     \"遇見\"\n"
"object:       \"David\"\n"
"location:     \"咖啡廳\"（可選）\n"
"time:         \"早上\"（可選）\n"
"strength:     0.9（0.0-1.0，越高越難忘）\n"
"access_count: 2（被讀取次數，影響衰減速度）\n"
"encoded_day:  3（第幾天存入）"
)

add_heading3(doc, "confusion.py — 困惑指數計算")
add_code_block(doc,
"C = w1 * U + w2 * K + w3 * S\n\n"
"U = 1.0 - (最高分行動 - 次高分行動) * 2   # 越難選擇，U 越高\n"
"K = 關鍵字衝突偵測（0 到 0.6）              # 如同時出現「陌生人」和「老朋友」\n"
"S = 場景新穎度（0.7）+ 行動停滯懲罰（0.2） # 一直重複同樣行動會被懲罰\n\n"
"每個角色的 w1、w2、w3 不同，體現個性差異。"
)

add_heading3(doc, "memory_consolidation.py — 睡眠濃縮 7 步驟")
steps = [
    "模型從今日 STM 命題中篩選重要的",
    "重要命題寫入 LTM",
    "模型生成 LTM 壓縮摘要（1-2 句話）",
    "對今天互動過的每個角色，模型更新關係描述",
    "模型從今天發生的事推斷當前情緒",
    "所有 LTM 命題依公式衰減，低分者刪除",
    "清空 STM，day + 1",
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    para_spacing(p, before=1, after=1)
    add_run(p, f"{i}. ", bold=True, color=C_ACCENT, size=10.5)
    add_run(p, s, color=C_GRAY, size=10.5)

# Model 層
add_heading2(doc, "Model 層（model/）— Phi-3.5-Vision 封裝")
model_modules = [
    ("model_loader.py", "自動選擇最佳裝置（CUDA > MPS > CPU）；Singleton 設計，整個系統只載入一次；提供 make_model_fn(max_new_tokens=N) 供 consolidation 使用"),
    ("text_encoder.py", "將 prompt 文字轉換為 Phi-3.5 的 chat 格式，並在有圖片時插入 <|image_1|> 等佔位符"),
    ("vision_encoder.py", "將 PIL 圖片轉換為模型所需的 pixel_values 張量格式"),
    ("fusion_decoder.py", "把文字和圖片 tensor 合併後呼叫 model.generate()，解碼後只回傳新生成的部分"),
    ("prompt_builder.py", "直覺路徑（系統說明+一句話個性+最近3筆STM+LTM摘要），思考路徑（完整個性+習慣+全部STM+相關LTM+完整關係摘要）"),
    ("output_parser.py", "用 regex 提取 [ACTION] / [THOUGHT] / [HAM] 三個區塊，將「前往:咖啡廳」拆成 verb 和 target"),
]
for name, desc in model_modules:
    add_heading3(doc, name)
    add_body(doc, desc)

# Agent 層
add_heading2(doc, "Agent 層（agent/）— 推論協調")
add_heading3(doc, "agent.py — 單角色一輪完整推論（8 步驟）")
agent_steps = ["計算困惑指數 → 決定推論路徑", "組裝 Prompt", "圖片前處理（有圖才跑）",
               "圖文融合", "模型推論", "解析輸出", "更新角色狀態（位置、行動）", "寫入 STM"]
for i, s in enumerate(agent_steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    para_spacing(p, before=1, after=1)
    add_run(p, f"{i}. ", bold=True, color=C_ACCENT, size=10.5)
    add_run(p, s, color=C_GRAY, size=10.5)

add_heading3(doc, "agent_manager.py — 5 角色協調器")
add_bullet(doc, "**AI 間對話傳遞**：當 A 說話對象是 B，manager 自動把 A 說的話傳給 B，讓 B 執行一輪推論並回應。用 _forwarded=True 旗標防止無限循環。")
add_bullet(doc, "**睡眠觸發**：偵測到 should_sleep=True 或 STM 滿 15 筆時，自動執行睡眠濃縮，濃縮完儲存到磁碟。全部角色都睡覺後，虛擬時鐘推進到下一天。")

# Perception / Server / World / Utils
for title, content in [
    ("Perception 層（perception/）— 視覺感知",
     "YOLO v8n（最輕量版本）偵測畫面後，只保留 10 類「有意義的物件」：人、椅子、杯子、筆電、瓶子、手提包、書、手機、長椅、桌子。"
     "偵測結果轉換成中文語意描述再注入 prompt，模型理解「2個人、1個杯子」比理解 bounding box 座標容易得多。"),
    ("Server 層（server/）— UE 通訊橋接",
     "WebSocket 伺服器接收 UE 傳來的訊息（character、scene、image_b64、input_text、target），"
     "圖片立刻做 Base64 解碼 → PIL → YOLO → 語意文字，再傳給 agent，最後回傳 action、thought、mode 給 UE。"),
    ("World 層（world/）— 虛擬時間",
     "虛擬時鐘以「午夜後分鐘數」儲存，每 tick 前進 30 分鐘（可設定），負責觸發固定排程（起床、吃飯等），"
     "所有角色都睡覺後才推進到隔天。"),
    ("Utils 層（utils/）— 工具",
     "啟動時全部載入記憶體，只有睡眠時才寫回磁碟——推論過程不需讀寫磁碟，速度快；一天內最多只損失一天的進度。"
     "每天產生一個 log 檔（logs/YYYYMMDD.log），記錄每輪的 C 值、行動決策、睡眠濃縮結果。"),
]:
    add_heading2(doc, title)
    add_body(doc, content)

add_divider(doc)

# ── 角色資料結構 ───────────────────────────────────────────────
add_heading1(doc, "角色資料結構（AI_Data/A_init.json）")
add_code_block(doc,
'{\n'
'  "name": "Amy",\n'
'  "name_code": "A",\n'
'  "role": "咖啡師",\n'
'  "personality_short": "（一句話版，直覺路徑 prompt 用）",\n'
'  "personality": "（完整版，思考路徑 prompt 用）",\n'
'  "habit": "每天早上先做一杯手沖咖啡...",\n'
'  "relationships": {\n'
'    "B": {\n'
'      "initial": "認識但關係微妙，對方似乎有好感",\n'
'      "summary": "（每次睡眠後由模型更新）"\n'
'    }\n'
'  },\n'
'  "emotion": "平靜",\n'
'  "stm": { "capacity": 15, "turns": [] },\n'
'  "ltm": { "ltm_summary": "", "propositions": [] },\n'
'  "state": { "day": 1, "current_location": "公寓", "current_action": "休息" },\n'
'  "schedule": { "slots": [ { "time": "07:00", "action": "起床", "type": "fixed" } ] },\n'
'  "confusion_weights": { "w1": 0.4, "w2": 0.3, "w3": 0.3, "threshold": 0.45 }\n'
'}'
)
add_body(doc, "initial 是角色「天生設定」，summary 是「活生生的記憶」，"
              "兩者共存讓角色既有基礎個性，又能因互動而改變對他人的看法。")
add_divider(doc)

# ── 模組結構 ──────────────────────────────────────────────────
add_heading1(doc, "模組結構")
add_code_block(doc,
"project/\n"
"├── AI_Data/                    # 角色 JSON（A-E 代號）\n"
"├── config/\n"
"│   ├── prompts.py              # 所有 prompt 模板（繁體中文）\n"
"│   ├── world_config.py         # 世界規則、情緒清單、衰減參數\n"
"│   ├── model_config.py         # 模型參數、token 預算\n"
"│   └── action_list.py          # 合法行動動詞與地點\n"
"├── core/                       # 純邏輯，不依賴模型，可獨立測試\n"
"│   ├── character.py\n"
"│   ├── stm.py\n"
"│   ├── ltm.py\n"
"│   ├── confusion.py\n"
"│   └── memory_consolidation.py\n"
"├── model/\n"
"│   ├── model_loader.py\n"
"│   ├── vision_encoder.py\n"
"│   ├── text_encoder.py\n"
"│   ├── fusion_decoder.py\n"
"│   ├── prompt_builder.py\n"
"│   └── output_parser.py\n"
"├── agent/\n"
"│   ├── agent.py\n"
"│   └── agent_manager.py\n"
"├── perception/\n"
"│   └── yolo_handler.py\n"
"├── server/\n"
"│   └── ws_server.py\n"
"├── world/\n"
"│   └── world_clock.py\n"
"├── utils/\n"
"│   ├── file_io.py\n"
"│   ├── logger.py\n"
"│   └── test_reset.py\n"
"├── logs/\n"
"├── simulate.py\n"
"├── test.ipynb\n"
"└── main.py"
)
add_divider(doc)

# ── 關鍵設計決策 ───────────────────────────────────────────────
add_heading1(doc, "關鍵設計決策")
add_styled_table(doc,
    headers=["決策", "為什麼這樣做"],
    rows=[
        ["HAM 三元組記憶（不用向量）", "不需要額外的 embedding 模型；命題可直接注入 context；人類也看得懂"],
        ["temperature = 0.0",          "輸出格式穩定，parser 不易出錯；同樣輸入可重現"],
        ["core 層不依賴 model 層",      "可用 mock 測試所有記憶邏輯，不需要 GPU"],
        ["YOLO 結果先語意化再給模型",   "模型理解「2個人」比理解座標數字容易得多"],
        ["狀態在 RAM，睡眠才存磁碟",    "推論速度快；一天內崩潰最多損失一天"],
        ["每步驟獨立 token 預算",       "情緒推斷只需 10 tokens，省 GPU 時間（5 角色睡眠從 47 分鐘降至 8 分鐘）"],
        ["對話轉發用 _forwarded 旗標",  "防止 A→B→A 無限對話迴圈"],
        ["兩路徑推論（直覺/思考）",     "模仿人腦的認知資源分配，平衡速度與品質"],
    ],
    col_widths=[5.5, 10.0],
)

# ── Save ─────────────────────────────────────────────────────
output_path = r"c:\Users\Rayyu\Desktop\agi\NEW\project\Project_description.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
